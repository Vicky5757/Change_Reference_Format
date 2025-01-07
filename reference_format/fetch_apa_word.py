import requests
import pandas as pd
from docx import Document

def fetch_metadata(title):
    """
    Fetch metadata for a given title using the CrossRef API.
    """
    url = "https://api.crossref.org/works"
    params = {"query.title": title, "rows": 1}
    try:
        response = requests.get(url, params=params, timeout=10)
        response.raise_for_status()
        data = response.json()

        if "message" in data and "items" in data["message"] and len(data["message"]["items"]) > 0:
            item = data["message"]["items"][0]

            # Process authors in APA format: Lastname, F. I.
            authors = []
            if "author" in item:
                for author in item["author"]:
                    given = author.get("given", "")
                    family = author.get("family", "")
                    if given and family:
                        initials = " ".join([name[0] + "." for name in given.split()])
                        authors.append(f"{family}, {initials}")

            metadata = {
                "Author": ", ".join(authors) if authors else "N/A",
                "Year": item.get("published-print", {}).get("date-parts", [[None]])[0][0] or "n.d.",
                "Title": title,
                "Journal": item.get("container-title", ["N/A"])[0] or "N/A",
                "Volume": item.get("volume", "N/A"),
                "Issue": item.get("issue", "N/A"),
                "Pages": item.get("page", "N/A"),
                "DOI": item.get("DOI", "N/A")
            }

            return metadata
        else:
            return {"Error": "No data found for this title."}
    except Exception as e:
        return {"Error": str(e)}

def format_apa(metadata):
    """
    Format metadata into APA 7th edition reference style.
    """
    if "Error" in metadata:
        return metadata["Error"]

    author = metadata["Author"]
    year = metadata["Year"]
    title = metadata["Title"]
    journal = metadata["Journal"]
    volume = metadata["Volume"]
    issue = f"({metadata['Issue']})" if metadata["Issue"] != "N/A" else ""
    pages = metadata["Pages"]
    doi = f"https://doi.org/{metadata['DOI']}" if metadata["DOI"] != "N/A" else "N/A"

    apa_reference = {
        "Author": author,
        "Year": year,
        "Title": title,
        "Journal": journal,
        "Volume": volume,
        "Issue": issue,
        "Pages": pages,
        "DOI": doi
    }

    return apa_reference

def save_to_word(references, output_word):
    """
    Save APA references to a Word document.
    """
    doc = Document()
    doc.add_heading("APA References", level=1)

    for ref in references:
        if isinstance(ref, dict) and "Author" in ref:
            para = doc.add_paragraph()
            para.add_run(f"{ref['Author']} ({ref['Year']}). {ref['Title']}. ").bold = False
            para.add_run(ref['Journal']).italic = True
            para.add_run(f", {ref['Volume']}{ref['Issue']}, {ref['Pages']}. {ref['DOI']}")
        else:
            doc.add_paragraph(ref)

    doc.save(output_word)

def process_titles_to_word(input_csv, output_word):
    """
    Process titles from an input CSV file and save APA references to a Word document.
    """
    try:
        df = pd.read_csv(input_csv)

        if "Title" not in df.columns:
            raise ValueError("Input CSV must contain a 'Title' column.")

        apa_references = []
        for title in df["Title"]:
            metadata = fetch_metadata(title)
            apa_reference = format_apa(metadata)
            apa_references.append(apa_reference)

        save_to_word(apa_references, output_word)
        print(f"APA formatted references saved to {output_word}")

    except FileNotFoundError as e:
        print(f"Error: {e}")
    except ValueError as e:
        print(f"Error: {e}")
    except Exception as e:
        print(f"Unexpected error: {e}")

if __name__ == "__main__":
    input_csv = "titles.csv"
    output_word = "apa_references.docx"
    process_titles_to_word(input_csv, output_word)
